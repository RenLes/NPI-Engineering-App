import io
from datetime import datetime

import pandas as pd


def export_word(title: str, sections: list[dict]) -> io.BytesIO:
    """Generate a branded Word document. sections = [{"heading": str, "body": str}, ...]"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("NETTS PLANNING AND INFRASTRUCTURE")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("Civil Engineering Platform")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph("")  # spacer

    # Title
    title_para = doc.add_heading(title, level=1)
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    # Meta
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {datetime.now().strftime('%d %B %Y %H:%M')}").font.size = Pt(9)

    doc.add_paragraph("")

    # Sections
    for sec in sections:
        h = doc.add_heading(sec.get("heading", ""), level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        doc.add_paragraph(sec.get("body", ""))

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer.add_run("Confidential — NPI Internal Use Only")
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def export_excel(title: str, dataframes: dict[str, pd.DataFrame]) -> io.BytesIO:
    """Generate an Excel workbook. dataframes = {"Sheet Name": DataFrame, ...}"""
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        for sheet_name, df in dataframes.items():
            df.to_excel(writer, sheet_name=sheet_name[:31], index=False)
    buf.seek(0)
    return buf


def _sanitize_pdf_text(text: str) -> str:
    """Replace Unicode characters unsupported by built-in PDF fonts."""
    replacements = {
        "\u2013": "-",   # en dash
        "\u2014": "-",   # em dash
        "\u2018": "'",   # left single quote
        "\u2019": "'",   # right single quote
        "\u201c": '"',   # left double quote
        "\u201d": '"',   # right double quote
        "\u2022": "*",   # bullet
        "\u2026": "...", # ellipsis
        "\u00b2": "2",   # superscript 2
        "\u00b3": "3",   # superscript 3
        "\u00b0": "deg", # degree
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def export_pdf(title: str, sections: list[dict]) -> io.BytesIO:
    """Generate a simple branded PDF. sections = [{"heading": str, "body": str}, ...]"""
    from fpdf import FPDF

    title = _sanitize_pdf_text(title)
    sections = [{"heading": _sanitize_pdf_text(s.get("heading", "")), "body": _sanitize_pdf_text(s.get("body", ""))} for s in sections]

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Header
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(27, 58, 92)
    pdf.cell(0, 12, "NETTS PLANNING AND INFRASTRUCTURE", ln=True, align="C")
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(107, 114, 128)
    pdf.cell(0, 6, "Civil Engineering Platform", ln=True, align="C")
    pdf.ln(8)

    # Title
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(27, 58, 92)
    pdf.cell(0, 10, title, ln=True)

    # Date
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(107, 114, 128)
    pdf.cell(0, 5, f"Generated: {datetime.now().strftime('%d %B %Y %H:%M')}", ln=True)
    pdf.ln(6)

    # Sections
    for sec in sections:
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(27, 58, 92)
        pdf.cell(0, 8, sec.get("heading", ""), ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(30, 30, 30)
        pdf.multi_cell(0, 5, sec.get("body", ""))
        pdf.ln(4)

    # Footer
    pdf.ln(10)
    pdf.set_font("Helvetica", "I", 7)
    pdf.set_text_color(156, 163, 175)
    pdf.cell(0, 5, "Confidential - NPI Internal Use Only", ln=True, align="C")

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf
